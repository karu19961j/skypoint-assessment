"""Resume text extraction + heuristic autofill suggestions.

Used by the upload endpoint to (a) populate `Application.resume_text` for
HR keyword search, and (b) suggest values the apply form can pre-fill.

Format support:
  - PDF  → pypdf.PdfReader (pure Python, no system deps)
  - DOCX → python-docx
  - DOC  → not extractable without `antiword`/`textract` (heavy native
           deps); upload is allowed but `extract_text` returns None.

Design choices:
  - Heuristics over heavy parsing. We don't try to LLM the resume; the
    autofill is intentionally narrow (skills + years-of-experience) so a
    bad guess doesn't bury a candidate with the wrong number.
  - Skill autofill is *job-aware*: the caller passes the JOB's required
    skills list and we return only those that the resume text mentions.
    That's reliable because we know the vocabulary in advance.
  - YOE extraction is best-effort regex; returns None on ambiguity.
"""

from __future__ import annotations

import io
import logging
import re
from dataclasses import dataclass

logger = logging.getLogger(__name__)


SUPPORTED_EXTENSIONS: frozenset[str] = frozenset({".pdf", ".doc", ".docx"})
SUPPORTED_CONTENT_TYPES: frozenset[str] = frozenset(
    {
        "application/pdf",
        "application/msword",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }
)


@dataclass
class AutofillSuggestion:
    """What the apply form pre-fills when a candidate uploads a resume.

    Empty fields mean "we couldn't tell" — the form leaves the user's
    current value alone in that case rather than wiping it.
    """

    skills: list[str]
    years_experience: int | None


def extension_for(filename: str) -> str:
    """Lowercase extension including the dot, e.g. ".pdf". Empty if none."""
    idx = filename.rfind(".")
    if idx < 0:
        return ""
    return filename[idx:].lower()


def extract_text(*, filename: str, body: bytes) -> str:
    """Return plain text extracted from a resume binary, or "" if we can't.

    Never raises — extraction failures are logged and treated as "no text".
    A resume that we can't parse is still uploadable and downloadable;
    it just won't be keyword-searchable or autofill-able.
    """
    ext = extension_for(filename)
    try:
        if ext == ".pdf":
            return _extract_pdf(body)
        if ext == ".docx":
            return _extract_docx(body)
        if ext == ".doc":
            # Binary .doc parsing needs antiword / olefile + heuristics;
            # not worth shipping the deps for. Document the gap in the
            # README and move on.
            return ""
    except Exception:  # pragma: no cover - logged for ops, not for tests
        logger.exception("Failed to extract text from resume %r", filename)
    return ""


def _extract_pdf(body: bytes) -> str:
    # Lazy import so test environments that don't exercise the upload
    # path don't pay the import cost.
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(body))
    chunks: list[str] = []
    for page in reader.pages:
        try:
            chunks.append(page.extract_text() or "")
        except Exception:
            # A single bad page shouldn't tank the whole extraction.
            continue
    return "\n".join(chunks).strip()


def _extract_docx(body: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(body))
    parts = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts).strip()


# ---------- autofill heuristics ----------


# Match "5+ years", "5 years", "5-7 years of experience", etc. We take
# the leading integer and (if a range) prefer the lower bound — biasing
# toward what the candidate can defend in an interview.
_YOE_PATTERNS: tuple[re.Pattern[str], ...] = (
    re.compile(r"(\d{1,2})\s*\+\s*year", re.IGNORECASE),
    re.compile(r"(\d{1,2})\s*-\s*\d{1,2}\s*year", re.IGNORECASE),
    re.compile(r"(\d{1,2})\s*year", re.IGNORECASE),
)


def suggest_autofill(
    *, resume_text: str, job_skills: list[str]
) -> AutofillSuggestion:
    """Return autofill candidates derived from `resume_text`.

    `job_skills` scopes the skill match to the vocabulary the job already
    declared — keeps the suggestion useful instead of dumping every
    technology mentioned anywhere.
    """
    if not resume_text:
        return AutofillSuggestion(skills=[], years_experience=None)

    text_lower = resume_text.lower()
    seen: set[str] = set()
    matched: list[str] = []
    for skill in job_skills:
        # Normalize once; preserve the job's casing in the suggestion so
        # the candidate's skill tags match the job's display.
        normalized = skill.strip().lower()
        if not normalized or normalized in seen:
            continue
        # Word-boundary match keeps "go" out of "google" and "r" out of
        # "ruby". `\b` is fine for ASCII tech skill names.
        pattern = re.compile(rf"\b{re.escape(normalized)}\b", re.IGNORECASE)
        if pattern.search(text_lower):
            matched.append(skill.strip())
            seen.add(normalized)

    return AutofillSuggestion(skills=matched, years_experience=_guess_yoe(text_lower))


def _guess_yoe(text_lower: str) -> int | None:
    """Best-effort YOE guess. Returns None when nothing plausible matched."""
    for pattern in _YOE_PATTERNS:
        match = pattern.search(text_lower)
        if not match:
            continue
        try:
            value = int(match.group(1))
        except (TypeError, ValueError):
            continue
        # Caps stay sane: a resume claiming "99 years experience" is
        # almost certainly OCR junk, not real data.
        if 0 <= value <= 60:
            return value
    return None
