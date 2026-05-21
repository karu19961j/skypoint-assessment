"""Tests for the resume upload / download / autofill pipeline.

Exercises:

  - text extraction (PDF + DOCX paths against in-memory bytes; .doc returns "")
  - autofill suggestion (skills cross-match against the job vocab, YOE regex)
  - the FastAPI upload endpoint (size cap + allow-list)
  - the download endpoint (owner check on both sides)
  - keyword search hitting resume_text (extension of the existing /by-job filter)

Real MinIO isn't needed — the `in_memory_storage` fixture in conftest.py
swaps the storage singleton for an in-process dict so the full route
flow runs end-to-end without a network dep.
"""

from __future__ import annotations

import io

import pytest
from fastapi.testclient import TestClient

from app.services.resume_text import (
    AutofillSuggestion,
    extract_text,
    suggest_autofill,
)

from .conftest import (
    InMemoryStorage,
    auth_headers,
    sample_application_payload,
    sample_job_payload,
)


# ---------- text extraction ----------


def test_extract_text_handles_unknown_extension_quietly():
    """No-format files return "" rather than raising — keeps upload OK."""
    assert extract_text(filename="resume.xyz", body=b"random bytes") == ""


def test_extract_text_doc_returns_empty():
    """Legacy .doc binary parsing isn't shipped; we still accept the upload
    but text extraction is a documented no-op for that format."""
    assert extract_text(filename="resume.doc", body=b"%doc-magic") == ""


def test_extract_text_docx_round_trips():
    """python-docx round-trips a paragraph and pulls it back out."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Five years of Python and FastAPI experience.")
    buf = io.BytesIO()
    doc.save(buf)
    out = extract_text(filename="resume.docx", body=buf.getvalue())
    assert "Python" in out
    assert "FastAPI" in out


# ---------- autofill suggestion ----------


def test_suggest_autofill_matches_only_job_skills():
    text = "Built apps in Python with Django and Postgres. Also know JavaScript and Go."
    sug = suggest_autofill(
        resume_text=text,
        job_skills=["Python", "FastAPI", "Postgres"],
    )
    # We only suggest skills the JOB asked for. Django + JS + Go aren't on
    # the wanted list, so they don't surface.
    assert set(sug.skills) == {"Python", "Postgres"}


def test_suggest_autofill_word_boundary_avoids_partial_matches():
    """Skill "go" should not match "google", "r" should not match "ruby"."""
    text = "Worked on google forms; ruby gems too."
    sug = suggest_autofill(resume_text=text, job_skills=["go", "r"])
    assert sug.skills == []


def test_suggest_autofill_extracts_years_of_experience():
    text = "I have 7 years of experience building distributed systems."
    sug = suggest_autofill(resume_text=text, job_skills=[])
    assert sug.years_experience == 7


def test_suggest_autofill_handles_missing_yoe():
    sug = suggest_autofill(resume_text="No yoe mentioned here.", job_skills=[])
    assert sug == AutofillSuggestion(skills=[], years_experience=None)


# ---------- upload endpoint ----------


def test_upload_rejects_unsupported_extension(
    client: TestClient,
    candidate_headers: dict[str, str],
    in_memory_storage: InMemoryStorage,
):
    resp = client.post(
        "/api/resume/upload",
        headers=candidate_headers,
        files={"file": ("resume.txt", b"not allowed", "text/plain")},
    )
    assert resp.status_code == 415
    assert "PDF" in resp.json()["detail"] or "pdf" in resp.json()["detail"].lower()


def test_upload_rejects_oversized_file(
    client: TestClient,
    candidate_headers: dict[str, str],
    in_memory_storage: InMemoryStorage,
):
    from app.config import get_settings

    cap = get_settings().resume_max_bytes
    # One byte over the cap.
    oversized = b"%PDF-1.4\n" + b"a" * cap
    resp = client.post(
        "/api/resume/upload",
        headers=candidate_headers,
        files={"file": ("resume.pdf", oversized, "application/pdf")},
    )
    assert resp.status_code == 413


def test_upload_stores_and_returns_key(
    client: TestClient,
    candidate_headers: dict[str, str],
    in_memory_storage: InMemoryStorage,
):
    """Successful upload returns the key shape the apply form needs."""
    resp = client.post(
        "/api/resume/upload",
        headers=candidate_headers,
        files={"file": ("resume.pdf", b"%PDF-1.4\ndummy-content", "application/pdf")},
    )
    assert resp.status_code == 201, resp.text
    body = resp.json()
    assert body["resume_key"].startswith("resumes/")
    assert body["resume_key"].endswith(".pdf")
    assert body["filename"] == "resume.pdf"
    assert body["size_bytes"] > 0
    # We stored the bytes — head_object returns metadata.
    assert in_memory_storage.head_object(body["resume_key"]) is not None


def test_upload_with_job_id_returns_skill_autofill(
    client: TestClient,
    hr_headers: dict[str, str],
    candidate_headers: dict[str, str],
    in_memory_storage: InMemoryStorage,
):
    """When job_id is passed, autofill cross-matches resume text against
    that job's required skills."""
    # Seed a job that wants Python + FastAPI
    job_resp = client.post(
        "/api/jobs/",
        headers=hr_headers,
        json=sample_job_payload(skills=["python", "fastapi"]),
    )
    job_id = job_resp.json()["id"]

    # Build a tiny DOCX with the matching skills in the text.
    from docx import Document

    doc = Document()
    doc.add_paragraph("Worked with python and fastapi for five years.")
    buf = io.BytesIO()
    doc.save(buf)

    resp = client.post(
        f"/api/resume/upload?job_id={job_id}",
        headers=candidate_headers,
        files={
            "file": (
                "resume.docx",
                buf.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert resp.status_code == 201
    auto = resp.json()["autofill"]
    assert set(auto["skills"]) == {"python", "fastapi"}
    assert auto["years_experience"] == 5


# ---------- download endpoint ----------


def test_download_streams_resume_for_owner(
    client: TestClient,
    hr_headers: dict[str, str],
    candidate_headers: dict[str, str],
    in_memory_storage: InMemoryStorage,
):
    # Seed a job + upload a resume + apply.
    job_id = client.post(
        "/api/jobs/", headers=hr_headers, json=sample_job_payload()
    ).json()["id"]

    upload = client.post(
        "/api/resume/upload",
        headers=candidate_headers,
        files={"file": ("cv.pdf", b"%PDF-1.4\ncontent", "application/pdf")},
    )
    resume_key = upload.json()["resume_key"]

    apply_resp = client.post(
        "/api/applications/",
        headers=candidate_headers,
        json=sample_application_payload(job_id, resume_key=resume_key),
    )
    assert apply_resp.status_code == 201, apply_resp.text
    application_id = apply_resp.json()["id"]

    # HR (owner of the job) can download.
    dl = client.get(f"/api/resume/{application_id}/download", headers=hr_headers)
    assert dl.status_code == 200
    assert dl.headers["content-type"] == "application/pdf"
    assert "cv.pdf" in dl.headers["content-disposition"]
    assert dl.content.startswith(b"%PDF-1.4")


def test_download_blocked_for_other_hr(
    client: TestClient,
    candidate_headers: dict[str, str],
    in_memory_storage: InMemoryStorage,
):
    """An HR who doesn't own the job can't download the resume."""
    from .conftest import register_user

    owner_hr = register_user(
        client, email="owner@example.com", password="HrPass1234!", role="hr",
        full_name="Owner HR",
    )
    other_hr = register_user(
        client, email="other@example.com", password="HrPass1234!", role="hr",
        full_name="Other HR",
    )
    job_id = client.post(
        "/api/jobs/", headers=auth_headers(owner_hr), json=sample_job_payload()
    ).json()["id"]
    upload = client.post(
        "/api/resume/upload",
        headers=candidate_headers,
        files={"file": ("cv.pdf", b"%PDF-1.4\ncontent", "application/pdf")},
    )
    application_id = client.post(
        "/api/applications/",
        headers=candidate_headers,
        json=sample_application_payload(job_id, resume_key=upload.json()["resume_key"]),
    ).json()["id"]

    dl = client.get(
        f"/api/resume/{application_id}/download",
        headers=auth_headers(other_hr),
    )
    assert dl.status_code == 403


def test_download_404_when_application_has_no_resume(
    client: TestClient,
    hr_headers: dict[str, str],
    candidate_headers: dict[str, str],
    in_memory_storage: InMemoryStorage,
):
    job_id = client.post(
        "/api/jobs/", headers=hr_headers, json=sample_job_payload()
    ).json()["id"]
    apply_resp = client.post(
        "/api/applications/",
        headers=candidate_headers,
        json=sample_application_payload(job_id, resume_key=None),
    )
    application_id = apply_resp.json()["id"]
    dl = client.get(f"/api/resume/{application_id}/download", headers=hr_headers)
    assert dl.status_code == 404


# ---------- keyword search picks up resume_text ----------


def test_applicant_keyword_search_finds_resume_text(
    client: TestClient,
    hr_headers: dict[str, str],
    candidate_headers: dict[str, str],
    in_memory_storage: InMemoryStorage,
):
    """HR keyword search ORs over resume_text — a candidate who has the
    keyword only in their resume (not skills/cover) still surfaces."""
    job_id = client.post(
        "/api/jobs/", headers=hr_headers, json=sample_job_payload()
    ).json()["id"]

    # DOCX with a unique-sounding word the HR can search for.
    from docx import Document

    doc = Document()
    doc.add_paragraph("Hands-on with kubernetes orchestration.")
    buf = io.BytesIO()
    doc.save(buf)
    upload = client.post(
        "/api/resume/upload",
        headers=candidate_headers,
        files={
            "file": (
                "cv.docx",
                buf.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    client.post(
        "/api/applications/",
        headers=candidate_headers,
        json=sample_application_payload(
            job_id,
            resume_key=upload.json()["resume_key"],
            cover_note="standard cover",
            skills=["python"],  # NOT kubernetes
        ),
    )

    # Search for "kubernetes" — should hit via resume_text even though
    # neither skills nor cover_note mention it.
    resp = client.get(
        f"/api/applications/by-job/{job_id}?q=kubernetes",
        headers=hr_headers,
    )
    assert resp.status_code == 200
    rows = resp.json()
    assert len(rows) == 1
